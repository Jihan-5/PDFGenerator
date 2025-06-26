from docx import Document
from pathlib import Path
import json
import logging
from PIL import Image
import io
import pytesseract

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %( sociala")
logger = logging.getLogger(__name__)

def parse_docx(path: str | Path) -> list[dict]:
    doc = Document(path)
    sections: list[dict] = []
    current = None
    para_iter = iter(doc.paragraphs)
    tables = iter(doc.tables)
    table_idx = 0
    images = []

    # Extract images for OCR
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            img_data = rel.target_part.blob
            images.append(Image.open(io.BytesIO(img_data)))

    def parse_table(table):
        fields = []
        header = [cell.text.strip().lower() for cell in table.rows[0].cells]
        start_row = 1 if any(h in {"field_name", "name", "column"} for h in header) else 0

        for row in table.rows[start_row:]:
            cells = [cell.text.strip() for cell in row.cells]
            if not cells or not cells[0]:
                logger.warning(f"Skipping empty row in table: {cells}")
                continue
            field = {
                "field_name": cells[0],
                "field_type": cells[1] if len(cells) > 1 else "",
                "description": cells[2] if len(cells) > 2 else ""
            }
            if field["field_name"] not in [f["field_name"] for f in fields]:
                fields.append(field)
                logger.debug(f"Extracted field: {field['field_name']}")
        return fields

    def parse_image_table(img: Image) -> list[dict]:
        # Basic OCR to extract table-like text
        text = pytesseract.image_to_string(img)
        fields = []
        for line in text.splitlines():
            parts = [p.strip() for p in line.split("|") if p.strip()]
            if len(parts) >= 2:
                field = {
                    "field_name": parts[0],
                    "field_type": parts[1] if len(parts) > 1 else "",
                    "description": parts[2] if len(parts) > 2 else ""
                }
                if field["field_name"] not in [f["field_name"] for f in fields]:
                    fields.append(field)
        return fields

    while True:
        try:
            para = next(para_iter)
        except StopIteration:
            break

        style = para.style.name
        text = para.text.strip()

        if style.startswith("Heading 1") and text:
            if current:
                sections.append(current)
            current = {"name": text.split(maxsplit=1)[-1], "description": "", "purpose": [], "key_fields": []}

        elif style.startswith("Heading 2") and text.lower().startswith("overview"):
            current["_mode"] = "desc"
        elif style.startswith("Heading 2") and text.lower().startswith("purpose"):
            current["_mode"] = "purpose"
        elif style.startswith("Heading 2") and text.lower().startswith("field"):
            current["_mode"] = "fields"
            if table_idx < len(doc.tables):
                table = doc.tables[table_idx]
                fields = parse_table(table)
                if not fields and images:  # Fallback to OCR if table is empty
                    logger.info(f"Table {current['name']} empty, trying OCR...")
                    fields = parse_image_table(images[table_idx % len(images)])
                if fields:
                    current["key_fields"].extend(fields)
                    logger.info(f"Extracted {len(fields)} fields from table {current['name']}")
                else:
                    logger.warning(f"No valid fields found in table {current['name']}")
                table_idx += 1
                current["_mode"] = None

        elif current and text:
            mode = current.get("_mode")
            if mode == "desc":
                current["description"] += " " + text
            elif mode == "purpose":
                current["purpose"].append(text.lstrip("• ").strip())
            elif mode == "fields":
                parts = [p.strip() for p in text.split("\t")]
                if len(parts) >= 3:
                    current["key_fields"].append(
                        {"field_name": parts[0], "field_type": parts[1], "description": parts[2]}
                    )

    if current:
        sections.append(current)
    for s in sections:
        s.pop("_mode", None)
    logger.info(f"Extracted {len(sections)} table specs from {path}")
    return sections